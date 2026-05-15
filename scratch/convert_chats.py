import os
import glob
from docx import Document
from docx.shared import Pt

# Mapping of known conversation IDs to titles
ID_TO_TITLE = {
    "252b637b-ea9b-4df6-9cc0-14792e0d219d": "Refining Nexa Role Permissions",
    "176b3a4d-ba60-4ae3-a5b9-faee49d739ea": "Refining Manager Permissions and Forms",
    "001c27d9-28ee-43f3-918c-65e2ad6a7b59": "Nexa Platform Stabilization",
    "96d7a325-8d60-48b1-8c77-cb146a82fe8e": "Refining Nexa Admin Interface",
    "41aef79f-2e0c-4218-9b50-f54b972ac723": "Implementing NS5 Design System",
    "4bd18138-7a6f-4c8b-8aee-485e6dcde207": "Implementing NS5 Design System",
    "2e489fd8-9b12-440e-9886-4def25155df7": "Applying Design System Standards",
    "4055d5f9-9da4-4ee3-b01d-e53f5bc95c3e": "Building Nexa System Admin Panel",
    "2852ba8c-3526-4322-ac63-8bd5e9a66bd4": "Optimizing Cloud Run Deployment",
    "60219c8d-e16b-4491-82fa-6b201e2a9e0e": "Deploying Nexa To Cloud Run"
}

BASE_DIR = "/home/snow/.gemini/antigravity/brain"
OUTPUT_DIR = os.path.expanduser("~/Documents/AntigravityChatHistory")

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

def create_docx(conv_id, content, title):
    doc = Document()
    doc.add_heading(title, 0)
    
    # Add content
    for line in content.split('\n'):
        p = doc.add_paragraph(line)
        # Optional: styling
        run = p.runs[0] if p.runs else p.add_run()
        run.font.size = Pt(10)

    # Clean filename
    safe_title = "".join([c if c.isalnum() else "_" for c in title])
    filename = f"{conv_id}_{safe_title}.docx"
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return path

def main():
    print(f"Scanning for conversations in {BASE_DIR}...")
    conv_dirs = glob.glob(os.path.join(BASE_DIR, "*"))
    
    for conv_dir in conv_dirs:
        if not os.path.isdir(conv_dir):
            continue
            
        conv_id = os.path.basename(conv_dir)
        log_path = os.path.join(conv_dir, ".system_generated", "logs", "overview.txt")
        
        if os.path.exists(log_path):
            print(f"Processing {conv_id}...")
            with open(log_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            title = ID_TO_TITLE.get(conv_id, f"Conversation_{conv_id}")
            save_path = create_docx(conv_id, content, title)
            print(f"Saved to {save_path}")

if __name__ == "__main__":
    main()
